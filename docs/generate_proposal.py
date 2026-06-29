"""
Run: python docs/generate_proposal.py
Produces: docs/DermAssist_AI_Proposal.docx
Requires: pip install python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────
title = doc.add_heading("DermAssist-AI", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Explainable Multimodal Skin Cancer Detection\nand Clinical Decision Support System")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True

doc.add_paragraph()

# ── Section helper ─────────────────────────────────────────────────────────
def section(heading, body_paras):
    doc.add_heading(heading, level=1)
    for p in body_paras:
        doc.add_paragraph(p, style="Normal")

# ── 1. Overview ────────────────────────────────────────────────────────────
section("1. Project Overview", [
    "DermAssist-AI is an AI-powered clinical decision support system that classifies common "
    "skin conditions — Acne, Eczema, Psoriasis, and Melanoma-like lesions — from photographs "
    "using deep convolutional neural networks with transfer learning. The system produces "
    "explainability overlays (Grad-CAM heatmaps) that highlight the image regions driving "
    "each prediction, making the model interpretable for both patients and clinicians.",
])

# ── 2. IEEE Areas ──────────────────────────────────────────────────────────
section("2. IEEE Technical Areas", [
    "• Computer Vision (image classification, object localisation)",
    "• Medical / Clinical AI (diagnostic decision support)",
    "• Deep Learning (CNN transfer learning, fine-tuning)",
    "• Explainable AI / XAI (Grad-CAM, LIME, SHAP)",
])

# ── 3. Problem Statement ───────────────────────────────────────────────────
section("3. Problem Statement", [
    "Skin conditions collectively affect over 1.9 billion people worldwide (WHO, 2023). "
    "Early detection of melanoma can increase the 5-year survival rate from ~20% (Stage IV) "
    "to ~99% (Stage I). However, access to certified dermatologists is limited in many "
    "regions. An AI-assisted pre-screening tool can bridge this gap, flag high-risk lesions, "
    "and direct patients to appropriate care faster.",
])

# ── 4. Objectives ─────────────────────────────────────────────────────────
section("4. Objectives", [
    "1. Build an image classification pipeline achieving ≥ 85% test accuracy across four "
    "skin condition classes.",
    "2. Generate Grad-CAM saliency maps to highlight suspicious lesion regions.",
    "3. Deploy an interactive Streamlit web application for real-time inference.",
    "4. Document model uncertainty and limitations in a clinically responsible manner.",
])

# ── 5. System Architecture ────────────────────────────────────────────────
section("5. System Architecture", [
    "Layer 1 — Input: JPEG/PNG skin image uploaded by the user.",
    "Layer 2 — Preprocessing: Resize to 224×224, normalise to ImageNet statistics, "
    "apply augmentation (flip, colour jitter) during training.",
    "Layer 3 — Feature Extraction: EfficientNet-B0 pre-trained on ImageNet, "
    "fine-tuned on the skin-lesion dataset.",
    "Layer 4 — Classifier Head: Fully-connected layer → Softmax → 4-class probability vector.",
    "Layer 5 — Explainability: Grad-CAM on the last convolutional block produces a "
    "class-activation heatmap overlaid on the original image.",
    "Layer 6 — UI: Streamlit frontend displays prediction, confidence bar chart, "
    "and Grad-CAM overlay.",
])

# ── 6. Dataset ────────────────────────────────────────────────────────────
section("6. Dataset", [
    "Primary: ISIC 2019 / 2020 Challenge Dataset (open-access, ≈ 33 000 dermoscopy images).",
    "Secondary: DermNet NZ, Kaggle Skin Disease Dataset for non-melanoma classes.",
    "Split: 70% train / 15% validation / 15% test (stratified by class).",
    "Augmentation: Random crops, horizontal/vertical flips, colour jitter, rotation.",
])

# ── 7. Tech Stack ─────────────────────────────────────────────────────────
section("7. Technology Stack", [
    "Language: Python 3.10+",
    "Deep Learning: PyTorch 2.x, torchvision (EfficientNet-B0)",
    "Computer Vision: OpenCV, Pillow, Albumentations",
    "Explainability: pytorch-grad-cam, LIME, SHAP",
    "Web Framework: Streamlit (primary), Flask (REST API option)",
    "Notebooks: Jupyter (EDA, training, evaluation)",
    "Version Control: Git + GitHub",
    "Optional MLOps: MLflow for experiment tracking",
])

# ── 8. Evaluation Metrics ─────────────────────────────────────────────────
section("8. Evaluation Metrics", [
    "• Accuracy, Precision, Recall, F1-Score (macro & per-class)",
    "• AUC-ROC curve per class",
    "• Confusion matrix",
    "• Qualitative assessment of Grad-CAM heatmap alignment with lesion boundaries",
])

# ── 9. Project Timeline ───────────────────────────────────────────────────
section("9. Project Timeline (8 weeks)", [
    "Week 1–2: Literature review, dataset collection, EDA",
    "Week 3–4: Model training, hyperparameter tuning",
    "Week 5:   Explainability integration (Grad-CAM, LIME)",
    "Week 6:   Streamlit UI development",
    "Week 7:   Testing, bias analysis, model cards",
    "Week 8:   Report writing, demo video, final submission",
])

# ── 10. Global / Societal Impact ──────────────────────────────────────────
section("10. Global and Societal Impact", [
    "The system is designed to be lightweight enough to run on low-end hardware or cloud "
    "free-tier, making it accessible in under-resourced healthcare settings across South Asia, "
    "Sub-Saharan Africa, and rural communities. It can be adapted to local dermatology datasets "
    "and integrated into mobile screening programmes.",
])

# ── 11. Ethical Considerations ────────────────────────────────────────────
section("11. Ethical Considerations", [
    "• The system is strictly a screening aid, not a diagnostic device.",
    "• All outputs include a mandatory disclaimer to consult a licensed dermatologist.",
    "• Dataset bias (skin tone distribution) is documented and mitigated via balanced sampling.",
    "• No personal health information is stored; images are discarded after inference.",
])

# ── 12. Team & Roles ──────────────────────────────────────────────────────
section("12. Team & Roles", [
    "• ML Engineer: Model architecture, training pipeline",
    "• Computer Vision Engineer: Preprocessing, augmentation, Grad-CAM",
    "• Full-Stack Developer: Streamlit UI, Flask API",
    "• Data Engineer: Dataset curation, preprocessing scripts",
])

doc.add_heading("References", level=1)
refs = [
    "Codella N. et al. (2019). Skin Lesion Analysis Toward Melanoma Detection 2018. arXiv:1902.03368.",
    "Selvaraju R. et al. (2017). Grad-CAM: Visual Explanations from Deep Networks. ICCV.",
    "Tan M. & Le Q. (2019). EfficientNet: Rethinking Model Scaling for CNNs. ICML.",
    "WHO (2023). Skin Diseases Fact Sheet.",
]
for r in refs:
    p = doc.add_paragraph(r, style="Normal")
    p.paragraph_format.left_indent = Pt(24)

doc.save("docs/DermAssist_AI_Proposal.docx")
print("Saved -> docs/DermAssist_AI_Proposal.docx")
